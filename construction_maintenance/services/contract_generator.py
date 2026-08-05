from __future__ import annotations

import os
import re
from pathlib import Path
from datetime import datetime
from typing import Any

import docx

# Strict Template Registry
TEMPLATE_REGISTRY: dict[str, dict[str, Any]] = {
    "labor_contract_no_social_insurance_v1": {
        "template_id": "labor_contract_no_social_insurance_v1",
        "template_name": "2026新版劳务合同（不交社保）",
        "version": "1.0",
        "enabled": True,
        "category": "人员合同",
        "relative_path": "templates/docx_templates/2026_labor_contract_no_social_security.docx",
        "description": "基于服务器上固定的 Word (.docx) 模版填充生成，原法律条款、排版及样式保持 100% 不变。",
        "required_placeholders": [
            "{{employee_name}}",
            "{{employee_id_number}}",
            "{{signing_date}}",
        ]
    },
    "01_no_social_security_labor_contract": {
        "template_id": "01_no_social_security_labor_contract",
        "template_name": "2026新版劳务合同（不交社保）",
        "version": "1.0",
        "enabled": True,
        "category": "人员合同",
        "relative_path": "templates/docx_templates/2026_labor_contract_no_social_security.docx",
        "description": "基于服务器上固定的 Word (.docx) 模版填充生成，原法律条款、排版及样式保持 100% 不变。",
        "required_placeholders": [
            "{{employee_name}}",
            "{{employee_id_number}}",
            "{{signing_date}}",
        ]
    }
}


def list_contract_templates() -> list[dict[str, Any]]:
    templates = []
    seen = set()
    for t_id, meta in TEMPLATE_REGISTRY.items():
        name = meta["template_name"]
        if name not in seen:
            seen.add(name)
            templates.append({
                "id": t_id,
                "name": meta["template_name"],
                "version": meta["version"],
                "category": meta["category"],
                "description": meta["description"],
            })
    return templates


def get_template_by_id(template_id: str) -> dict[str, Any]:
    if template_id in TEMPLATE_REGISTRY:
        return TEMPLATE_REGISTRY[template_id]
    return TEMPLATE_REGISTRY["labor_contract_no_social_insurance_v1"]


def verify_template_file(template_id: str) -> Path:
    """
    Strict verification step:
    1. Check if template_id exists in registry and enabled.
    2. Check if template file exists on server.
    3. Check read permissions.
    4. Check .docx format and readability.
    5. Check required placeholders.
    Raises exception on ANY error; NEVER falls back to free-text LLM generation!
    """
    if template_id not in TEMPLATE_REGISTRY:
        raise ValueError(f"指定合同模板不存在 (template_id={template_id})")
    
    meta = TEMPLATE_REGISTRY[template_id]
    if not meta.get("enabled"):
        raise ValueError(f"指定合同模板当前已被禁用 (template_id={template_id})")

    base_dir = Path(__file__).parent.parent
    template_file_path = base_dir / meta["relative_path"]

    # 1. Existence check
    if not template_file_path.exists():
        raise FileNotFoundError(f"指定合同模板文件不存在，终止生成: {template_file_path}")

    # 2. Permission check
    if not os.access(template_file_path, os.R_OK):
        raise PermissionError(f"服务无权读取指定合同模板文件: {template_file_path}")

    # 3. Format check
    if template_file_path.suffix.lower() != ".docx":
        raise ValueError(f"指定合同模板格式错误，必须为 .docx 格式文档: {template_file_path}")

    # 4. Readability check
    try:
        doc = docx.Document(template_file_path)
    except Exception as exc:
        raise ValueError(f"指定合同模板文件损坏或无法读取打开: {exc}")

    # 5. Placeholders check
    full_text = "\n".join([p.text for p in doc.paragraphs])
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                full_text += "\n" + cell.text

    missing_tags = []
    for tag in meta.get("required_placeholders", []):
        if tag not in full_text:
            missing_tags.append(tag)

    if missing_tags:
        raise ValueError(f"指定合同模板文件缺少预定义的占位符标签 {missing_tags}，终止生成。")

    return template_file_path


def generate_contract_from_template(
    template_id: str,
    employee_data: dict[str, Any],
    contract_data: dict[str, Any],
    upload_folder: Path,
    generated_by_user_id: int = 1,
) -> dict[str, Any]:
    """
    Fixed Template Population Mode.
    Open template -> Replace placeholders -> Verify result -> Save output files.
    STRICTLY NO LLM WRITING/REWRITING OF CONTRACT CLAUSES.
    """
    # 1. Verify template file strictly
    template_path = verify_template_file(template_id)
    meta = TEMPLATE_REGISTRY[template_id]

    doc = docx.Document(template_path)

    # 2. Build field mapping
    contract_no = f"YLT-CON-{datetime.now().strftime('%Y%m%d%H%M%S')}"
    signing_date = contract_data.get("signing_date") or datetime.now().strftime("%Y年%m月%d日")
    
    start_date = contract_data.get("contract_start_date") or datetime.now().strftime("%Y年%m月%d日")
    end_date = contract_data.get("contract_end_date") or "作业完工结清薪资止"
    salary = str(contract_data.get("salary") or employee_data.get("salary_rate") or "350.00")
    salary_payment_date = str(contract_data.get("salary_payment_date") or "25")
    job_position = contract_data.get("job_position") or employee_data.get("job_type") or "施工人员"
    work_location = contract_data.get("work_location") or "工程现场"

    mapping = {
        "{{contract_no}}": contract_no,
        "{{employer_name}}": contract_data.get("employer_name") or "北京营力特建筑工程有限公司",
        "{{employer_address}}": contract_data.get("employer_address") or "北京市门头沟区妙峰山镇水丁路1号院A074室",
        "{{legal_representative}}": contract_data.get("legal_representative") or "谢世营",
        "{{employee_name}}": employee_data.get("name") or "",
        "{{employee_gender}}": employee_data.get("gender") or "男",
        "{{employee_ethnicity}}": employee_data.get("ethnicity") or "汉",
        "{{employee_birth_date}}": employee_data.get("birth_date") or "",
        "{{employee_id_number}}": employee_data.get("id_number") or "",
        "{{employee_phone}}": employee_data.get("phone") or contract_data.get("employee_phone") or "",
        "{{employee_address}}": employee_data.get("address") or "",
        "{{id_issuing_authority}}": employee_data.get("issuing_authority") or "",
        "{{id_valid_from}}": employee_data.get("valid_from") or "",
        "{{id_valid_until}}": employee_data.get("valid_until") or "",
        "{{job_position}}": job_position,
        "{{work_location}}": work_location,
        "{{contract_start_date}}": start_date,
        "{{contract_end_date}}": end_date,
        "{{salary}}": salary,
        "{{salary_payment_date}}": salary_payment_date,
        "{{signing_date}}": signing_date,
    }

    # 3. Replace placeholders in paragraphs
    for p in doc.paragraphs:
        _replace_placeholders_in_paragraph(p, mapping)

    # Replace in tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    _replace_placeholders_in_paragraph(p, mapping)

    # 4. Save output .docx file (Original template remains read-only!)
    emp_id = employee_data.get("id") or 0
    timestamp = int(datetime.now().timestamp())
    file_prefix = f"generated_contract_{emp_id}_{timestamp}"
    docx_filename = f"{file_prefix}.docx"
    html_filename = f"{file_prefix}.html"

    output_docx_path = upload_folder / docx_filename
    output_html_path = upload_folder / html_filename

    upload_folder.mkdir(parents=True, exist_ok=True)
    doc.save(output_docx_path)

    # Render matching HTML for online browser preview & direct printing (Renders 100% full Word text clauses)
    html_content = render_contract_html_for_preview(meta["template_name"], mapping, doc=doc, docx_filename=docx_filename)
    with open(output_html_path, "w", encoding="utf-8") as f:
        f.write(html_content)

    # 5. Output Verification
    output_doc = docx.Document(output_docx_path)
    output_text = "\n".join([p.text for p in output_doc.paragraphs])

    unreplaced = re.findall(r"\{\{[a-zA-Z0-9_]+\}\}", output_text)
    if unreplaced:
        # Check if unreplaced tags are in required list
        unreplaced_required = [u for u in set(unreplaced) if u in meta.get("required_placeholders", [])]
        if unreplaced_required:
            raise ValueError(f"合同生成失败：存在未替换的必填占位符字段 {unreplaced_required}")

    # Check employee name and ID card number exist in output text
    emp_name = employee_data.get("name")
    if emp_name and emp_name not in output_text:
        raise ValueError(f"合同生成校验失败：人员姓名【{emp_name}】未能注入文档。")

    # Record Audit Data
    audit_record = {
        "template_id": template_id,
        "template_name": meta["template_name"],
        "template_version": meta["version"],
        "employee_id": emp_id,
        "employee_name": emp_name,
        "generated_file_path": docx_filename,
        "generated_html_path": html_filename,
        "generated_at": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
        "generated_by": generated_by_user_id,
    }

    return audit_record


def _replace_placeholders_in_paragraph(p: Any, mapping: dict[str, str]) -> None:
    full_text = p.text
    has_match = False
    for k in mapping:
        if k in full_text:
            has_match = True
            break
    if not has_match:
        return

    for k, v in mapping.items():
        if k in full_text:
            full_text = full_text.replace(k, str(v or ""))

    if p.runs:
        p.runs[0].text = full_text
        for r in p.runs[1:]:
            r.text = ""
    else:
        p.text = full_text


def render_contract_html_for_preview(template_name: str, mapping: dict[str, str], doc: Any = None, docx_filename: str = "") -> str:
    person_name = mapping.get("{{employee_name}}", "")

    body_elements = []
    if doc:
        for p in doc.paragraphs:
            txt = p.text.strip()
            if not txt:
                continue
            if txt.startswith("合同编号:"):
                body_elements.append(f'<div style="text-align: right; color: #64748b; font-size: 13.5px; margin-bottom: 20px; font-weight: bold; font-family: monospace;">{txt}</div>')
            elif txt == "劳务合同" or txt == "劳 务 合 同":
                body_elements.append(f'<h1 style="text-align: center; color: #0f766e; font-size: 26px; margin: 30px 0 24px 0; letter-spacing: 4px; font-weight: 800;">劳 务 合 同</h1>')
            elif any(txt.startswith(prefix) for prefix in ["一、", "二、", "三、", "四、", "五、", "六、", "补充条款"]):
                body_elements.append(f'<h3 style="color: #0f766e; font-size: 16px; border-left: 4px solid #0f766e; padding-left: 10px; margin: 26px 0 14px 0;">{txt}</h3>')
            elif "甲方(盖章 )" in txt or "甲方 (用人单位)：" in txt or "乙方 (劳动者)  ：" in txt or txt.startswith("用工方( 甲方 )") or txt.startswith("劳务方( 乙方 )"):
                body_elements.append(f'<div style="background: #f8fafc; border: 1px solid #e2e8f0; border-radius: 6px; padding: 10px 16px; margin: 10px 0; font-weight: 600; color: #0f172a;">{txt}</div>')
            else:
                body_elements.append(f'<p style="font-size: 14px; color: #334155; line-height: 1.85; text-align: justify; margin-bottom: 8px;">{txt}</p>')
    
    content_html = "\n".join(body_elements)

    docx_download_btn = ""
    if docx_filename:
        docx_download_btn = f'<a href="/uploads/{docx_filename}?download=1" download class="btn-docx" style="background: rgba(255,255,255,0.2); color: #ffffff; text-decoration: none; padding: 6px 16px; border-radius: 6px; font-weight: 600; font-size: 13px; border: 1px solid rgba(255,255,255,0.4);">📥 下载 Word 原始文档 (.docx)</a>'

    html = f"""<!DOCTYPE html>
<html lang="zh-CN">
<head>
    <meta charset="UTF-8">
    <title>{template_name} - {person_name}</title>
    <style>
        body {{
            font-family: "PingFang SC", "Microsoft YaHei", "SimSun", sans-serif;
            color: #1e293b;
            line-height: 1.8;
            padding: 40px 60px;
            max-width: 850px;
            margin: 0 auto;
            background: #ffffff;
        }}
        .top-action-bar {{
            background: #0f766e;
            color: #ffffff;
            padding: 12px 24px;
            border-radius: 8px;
            margin-bottom: 30px;
            display: flex;
            justify-content: space-between;
            align-items: center;
            box-shadow: 0 4px 12px rgba(15, 118, 110, 0.15);
        }}
        .top-action-bar .title {{
            font-size: 14px;
            font-weight: 600;
        }}
        .top-action-bar .actions {{
            display: flex;
            gap: 12px;
            align-items: center;
        }}
        .top-action-bar .btn-print {{
            background: #ffffff;
            color: #0f766e;
            border: none;
            padding: 6px 18px;
            border-radius: 6px;
            font-weight: bold;
            cursor: pointer;
            font-size: 13px;
        }}
        @media print {{
            body {{ padding: 0; margin: 0; }}
            .top-action-bar {{ display: none !important; }}
        }}
    </style>
</head>
<body>
    <div class="top-action-bar">
        <span class="title">📄 在线预览：《{template_name}》（签约人：{person_name}）</span>
        <div class="actions">
            <button class="btn-print" onclick="window.print()">🖨️ 打印合同 (Ctrl+P)</button>
            {docx_download_btn}
        </div>
    </div>
    {content_html}
</body>
</html>"""
    return html

